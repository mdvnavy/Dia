import re
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def markdown_to_txt(markdown_text: str) -> str:
    """Converts markdown text into a clean plain text format."""
    lines = markdown_text.splitlines()
    output_lines = []
    for line in lines:
        match_h = re.match(r'^(#{1,6})\s+(.*)$', line)
        if match_h:
            level = len(match_h.group(1))
            text = match_h.group(2).strip()
            if level == 1:
                output_lines.append("")
                output_lines.append(text.upper())
                output_lines.append("=" * len(text))
            elif level == 2:
                output_lines.append("")
                output_lines.append(text)
                output_lines.append("-" * len(text))
            else:
                output_lines.append("")
                output_lines.append(f"=== {text} ===")
        elif line.strip() in ("---", "***"):
            output_lines.append("")
            output_lines.append("-" * 40)
            output_lines.append("")
        else:
            output_lines.append(line)
    return "\n".join(output_lines)

def _add_formatted_text_to_paragraph(paragraph, line: str):
    """Parses simple inline markdown (**bold**, *italic*) and adds runs to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def _build_docx_table(doc, rows: list[str]):
    """Parses markdown table rows and inserts a styled DOCX table."""
    parsed_rows = []
    for r in rows:
        cells = [c.strip() for c in r.strip('|').split('|')]
        # Skip separator line (e.g. |---|---|)
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return
        
    num_cols = max(len(row) for row in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Light Shading Accent 1'
    
    for row_idx, cells in enumerate(parsed_rows):
        row = table.rows[row_idx]
        for col_idx, cell_val in enumerate(cells):
            if col_idx < len(row.cells):
                # Clean up potential bold styling markup in docx cell text
                clean_val = re.sub(r'\*\*|__', '', cell_val)
                row.cells[col_idx].text = clean_val

def markdown_to_docx(markdown_text: str) -> bytes:
    """Converts markdown text to Word Document (.docx) bytes."""
    doc = Document()
    
    # Premium Typography Settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = markdown_text.splitlines()
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Page Breaks / Horizontal Rules
        if line.strip() in ("---", "***"):
            if in_table:
                _build_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            doc.add_page_break()
            i += 1
            continue
            
        # Headings
        match_h = re.match(r'^(#{1,6})\s+(.*)$', line)
        if match_h:
            if in_table:
                _build_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            level = len(match_h.group(1))
            text = match_h.group(2).strip()
            heading = doc.add_heading(text, level=level)
            heading.style.font.name = 'Arial'
            i += 1
            continue
            
        # Bullet Lists
        match_bullet = re.match(r'^[-*+]\s+(.*)$', line)
        if match_bullet:
            if in_table:
                _build_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text_to_paragraph(p, match_bullet.group(1).strip())
            i += 1
            continue
            
        # Numbered Lists
        match_num = re.match(r'^\d+\.\s+(.*)$', line)
        if match_num:
            if in_table:
                _build_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text_to_paragraph(p, match_num.group(1).strip())
            i += 1
            continue
            
        # Table Processing
        if line.strip().startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            if not line.strip():
                _build_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            else:
                table_rows.append(line)
            i += 1
            continue
            
        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph()
            _add_formatted_text_to_paragraph(p, line.strip())
        
        i += 1
        
    if in_table:
        _build_docx_table(doc, table_rows)
        
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def _markdown_inline_to_html(text: str) -> str:
    """Escapes HTML special chars and maps markdown bold/italic tags for reportlab Paragraphs."""
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    return text

def _build_pdf_table(rows: list[str], body_style: ParagraphStyle) -> Table:
    """Parses markdown table rows into a styled ReportLab Table flowable with wrapped cells."""
    parsed_rows = []
    for r in rows:
        cells = [c.strip() for c in r.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return Spacer(1, 1)
        
    formatted_data = []
    for row_idx, cells in enumerate(parsed_rows):
        formatted_row = []
        for cell_val in cells:
            cell_html = _markdown_inline_to_html(cell_val)
            if row_idx == 0:
                cell_html = f"<b>{cell_html}</b>"
            formatted_row.append(Paragraph(cell_html, body_style))
        formatted_data.append(formatted_row)
        
    col_count = max(len(r) for r in formatted_data)
    available_width = 504 # Letter page width 612 - 108 (left+right margins)
    col_width = available_width / col_count
    
    t = Table(formatted_data, colWidths=[col_width] * col_count)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0ede6')), # Light header background
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#16150f')),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d9d4c9')),
        ('TOPPADDING', (0,1), (-1,-1), 6),
        ('BOTTOMPADDING', (0,1), (-1,-1), 6),
    ]))
    return t

def markdown_to_pdf(markdown_text: str) -> bytes:
    """Converts markdown text to PDF bytes using ReportLab SimpleDocTemplate."""
    styles = getSampleStyleSheet()
    
    normal_style = styles['Normal']
    normal_style.textColor = colors.HexColor('#2e2c27')
    normal_style.fontSize = 10
    normal_style.leading = 14
    
    body_style = ParagraphStyle(
        'PDFBody',
        parent=normal_style,
        spaceAfter=8
    )
    
    title_style = ParagraphStyle(
        'PDFTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#75591f'),
        spaceBefore=14,
        spaceAfter=12,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'PDFH2',
        parent=styles['Heading2'],
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#5c4517'),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    
    h3_style = ParagraphStyle(
        'PDFH3',
        parent=styles['Heading3'],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#16150f'),
        spaceBefore=8,
        spaceAfter=4,
        keepWithNext=True
    )
    
    bullet_style = ParagraphStyle(
        'PDFBullet',
        parent=body_style,
        leftIndent=20,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    story = []
    lines = markdown_text.splitlines()
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Page breaks
        if line.strip() in ("---", "***"):
            if in_table:
                story.append(_build_pdf_table(table_rows, body_style))
                in_table = False
                table_rows = []
            story.append(PageBreak())
            i += 1
            continue
            
        # Headings
        match_h = re.match(r'^(#{1,3})\s+(.*)$', line)
        if match_h:
            if in_table:
                story.append(_build_pdf_table(table_rows, body_style))
                in_table = False
                table_rows = []
            level = len(match_h.group(1))
            text = match_h.group(2).strip()
            style = title_style if level == 1 else (h2_style if level == 2 else h3_style)
            story.append(Paragraph(_markdown_inline_to_html(text), style))
            i += 1
            continue
            
        # Bullet Lists
        match_bullet = re.match(r'^[-*+]\s+(.*)$', line)
        if match_bullet:
            if in_table:
                story.append(_build_pdf_table(table_rows, body_style))
                in_table = False
                table_rows = []
            text = match_bullet.group(1).strip()
            story.append(Paragraph(f"&bull; {_markdown_inline_to_html(text)}", bullet_style))
            i += 1
            continue
            
        # Numbered Lists
        match_num = re.match(r'^(\d+)\.\s+(.*)$', line)
        if match_num:
            if in_table:
                story.append(_build_pdf_table(table_rows, body_style))
                in_table = False
                table_rows = []
            num = match_num.group(1)
            text = match_num.group(2).strip()
            story.append(Paragraph(f"{num}. {_markdown_inline_to_html(text)}", bullet_style))
            i += 1
            continue
            
        # Table detection
        if line.strip().startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            if not line.strip():
                story.append(_build_pdf_table(table_rows, body_style))
                in_table = False
                table_rows = []
            else:
                table_rows.append(line)
            i += 1
            continue
            
        if line.strip():
            story.append(Paragraph(_markdown_inline_to_html(line.strip()), body_style))
        i += 1
        
    if in_table:
        story.append(_build_pdf_table(table_rows, body_style))
        
    out = BytesIO()
    doc = SimpleDocTemplate(
        out,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    doc.build(story)
    return out.getvalue()
